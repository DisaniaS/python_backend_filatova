from docx import Document
from fastapi.params import Depends

from report_data.repository import ReportDataRepository
from report_data.schema import ReportDataCreate


class ReportDataService:
    def __init__(self,
                 report_data_repo: ReportDataRepository = Depends()):
        self.report_data_repo = report_data_repo

    def get_report_data_by_report_id(self, report_id: int):
        db_report_data = self.report_data_repo.get_by_report_id(report_id)
        return db_report_data

    async def create_report_data(self, path: str, report_id: int):
        doc = Document(path)
        data = {}

        # Обработка таблиц (предполагаем структуру: первый столбец - ключ, второй - значение)
        for table in doc.tables:
            for row in table.rows:
                if len(row.cells) >= 2:  # Проверяем, что есть как минимум 2 столбца
                    key = row.cells[0].text.strip()
                    value = row.cells[1].text.strip()
                    self._process_key_value(key, value, data)

        # Обработка параграфов (формат "ключ: значение")
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if ":" in text:
                key, value = text.split(":", 1)
                self._process_key_value(key.strip(), value.strip(), data)

        self.report_data_repo.create(ReportDataCreate(
            report_id=report_id,
            system_name=data.get("system_name"),
            test_date=data.get("test_date"),
            department=data.get("department"),
            system_type=data.get("system_type"),
            test_time=data.get("test_time"),
            system_number=data.get("system_number"),
            latitude=data.get("latitude"),
            azimuth_minus_50=data.get("azimuth_minus_50"),
            azimuth_plus_50=data.get("azimuth_plus_50"),
            azimuth_nku=data.get("azimuth_nku"),
            repeated_azimuth_minus_50=data.get("repeated_azimuth_minus_50"),
            repeated_azimuth_plus_50=data.get("repeated_azimuth_plus_50"),
            repeated_azimuth_nku=data.get("repeated_azimuth_nku"),
            azimuth_determination_time=data.get("azimuth_determination_time"),
            table_position_exact=data.get("table_position_exact"),
            table_position_repeated=data.get("table_position_repeated"),
            humidity=data.get("humidity"),
            vibration_level=data.get("vibration_level"),
            calculated=False
        ))

    def _process_key_value(self, key: str, value: str, data: dict):
        try:
            value = value.replace(",", ".")

            if "Результаты испытаний системы" in key:
                data["system_name"] = str(value)
            elif "Дата проверки системы" in key:
                data["test_date"] = str(value)
            elif "Часть" in key:
                data["department"] = str(value)
            elif "Тип" in key:
                data["system_type"] = str(value)
            elif "Время испытания" in key:
                data["test_time"] = float(value)
            elif "Номер системы" in key:
                data["system_number"] = int(value)
            elif "Широта местоположения" in key:
                data["latitude"] = float(value)
            elif "Точное значение азимута при t = «-50 С»" in key:
                data["azimuth_minus_50"] = float(value)
            elif "Точное значение азимута при t = «+50 С»" in key:
                data["azimuth_plus_50"] = float(value)
            elif "Точное значение азимута в НКУ" in key:
                data["azimuth_nku"] = float(value)
            elif "Повторное значение азимута при t = «-50 С»" in key:
                data["repeated_azimuth_minus_50"] = float(value)
            elif "Повторное значение азимута при t = «+50 С»" in key:
                data["repeated_azimuth_plus_50"] = float(value)
            elif "Повторное значение азимута в НКУ" in key:
                data["repeated_azimuth_nku"] = float(value)
            elif "Время определения точного и повторного азимута" in key:
                data["azimuth_determination_time"] = float(value)
            elif "Положение стола для определения точного азимута" in key:
                data["table_position_exact"] = float(value)
            elif "Положение стола для определения повторного азимута" in key:
                data["table_position_repeated"] = float(value)
            elif "Влажность" in key:
                data["humidity"] = float(value)
            elif "Уровень вибрации" in key:
                data["vibration_level"] = float(value)
        except (ValueError, AttributeError):
            pass